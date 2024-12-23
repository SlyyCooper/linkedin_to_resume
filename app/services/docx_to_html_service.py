from docx import Document
from bs4 import BeautifulSoup
import os
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import time
import re
from typing import Optional

class DocxToHtmlConverter:
    """Converts DOCX files to styled HTML."""
    
    @staticmethod
    def convert_docx_to_html(docx_path: str, output_html_path: str) -> Optional[str]:
        """
        Convert a DOCX file to HTML with styling preserved.
        Returns the HTML content if successful, None if failed.
        """
        try:
            doc = Document(docx_path)
            
            # Start building HTML content (without full HTML document wrapper)
            html = []
            
            # Process each paragraph
            for para in doc.paragraphs:
                if not para.text.strip():  # Skip empty paragraphs
                    continue
                    
                # Determine style and format accordingly
                style_name = para.style.name.lower() if para.style else ''
                
                if 'heading 1' in style_name:
                    html.append(f'<h1 class="profile-name">{para.text}</h1>')
                elif 'heading 2' in style_name:
                    html.append(f'<h2 class="profile-headline">{para.text}</h2>')
                elif 'heading 3' in style_name:
                    html.append(f'<h3>{para.text}</h3>')
                elif 'list' in style_name:
                    # Handle list items
                    if not html[-1].startswith('<ul>'):
                        html.append('<ul class="skills-list">')
                    html.append(f'<li class="skill-tag">{para.text}</li>')
                    if not doc.paragraphs[doc.paragraphs.index(para) + 1:] or 'list' not in doc.paragraphs[doc.paragraphs.index(para) + 1].style.name.lower():
                        html.append('</ul>')
                else:
                    # Regular paragraph
                    text = para.text
                    
                    # Handle bold and italic
                    for run in para.runs:
                        if run.bold:
                            text = text.replace(run.text, f'<strong>{run.text}</strong>')
                        if run.italic:
                            text = text.replace(run.text, f'<em>{run.text}</em>')
                    
                    html.append(f'<p class="profile-about">{text}</p>')
            
            # Write the HTML file
            html_content = '\n'.join(html)
            with open(output_html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return html_content
            
        except Exception as e:
            print(f"Error converting DOCX to HTML: {str(e)}")
            return None

class DocxWatcher(FileSystemEventHandler):
    """Watches for changes in DOCX files and converts them to HTML."""
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        self.converter = DocxToHtmlConverter()
        
    def on_modified(self, event):
        if not event.is_directory and event.src_path.endswith('.docx'):
            print(f"DOCX file modified: {event.src_path}")
            base_name = os.path.splitext(os.path.basename(event.src_path))[0]
            html_path = os.path.join(self.output_dir, f"{base_name}.html")
            
            # Convert DOCX to HTML
            if self.converter.convert_docx_to_html(event.src_path, html_path):
                print(f"Successfully converted {event.src_path} to {html_path}")
            else:
                print(f"Failed to convert {event.src_path}")

def start_watcher(output_dir: str):
    """Start watching the output directory for DOCX changes."""
    event_handler = DocxWatcher(output_dir)
    observer = Observer()
    observer.schedule(event_handler, output_dir, recursive=False)
    observer.start()
    
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()

# Initial conversion of existing DOCX files
def convert_existing_files(output_dir: str):
    """Convert any existing DOCX files in the output directory."""
    converter = DocxToHtmlConverter()
    for filename in os.listdir(output_dir):
        if filename.endswith('.docx'):
            docx_path = os.path.join(output_dir, filename)
            html_path = os.path.join(output_dir, f"{os.path.splitext(filename)[0]}.html")
            if converter.convert_docx_to_html(docx_path, html_path):
                print(f"Successfully converted existing file {docx_path} to {html_path}")

if __name__ == "__main__":
    output_dir = "output"
    convert_existing_files(output_dir)
    start_watcher(output_dir)
