#!/usr/bin/env python3
"""
LinkedIn Profile "Highlight & Copy" Extractor
--------------------------------------------
This script uses Selenium (Python) to:
1. Prompt for LinkedIn credentials (username/email & password).
2. Prompt for the LinkedIn profile URL.
3. Log in via LinkedIn's login page.
   - If a verification puzzle appears, it pauses, allows the user to solve it manually,
     and continues once the user presses Enter.
4. Navigate to the specified profile URL.
5. Locate & click "see more" buttons to expand hidden sections.
6. Programmatically "highlight everything" on the page (simulating Ctrl + A).
7. Extract all visible text from the entire page post-expansion.
8. Save it into an 'output/profile.marathon' file.

It also includes optional logic to parse the raw text into a structured format
using GPT-4o. If you want just the raw text extraction, you can remove or ignore
the "structure_profile_data" and "save_structured_profile" parts.

DISCLAIMER:
- This script is a proof-of-concept and may violate LinkedIn's User Agreement
  if used in production. Use responsibly and lawfully.
- LinkedIn may change its HTML structure, so selectors may need to be updated.
"""

import os
import time
import getpass
from typing import List, Optional, Dict, Any

# Selenium imports
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# Optional: GPT-based parsing (requires your environment set up accordingly)
from openai import OpenAI
from pydantic import BaseModel
from dotenv import load_dotenv

# Optional: For Word/Markdown conversions
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

load_dotenv()

# ------------------------------
# 1. Data Models (Optional GPT)
# ------------------------------
class Experience(BaseModel):
    title: str
    company: str
    duration: str
    description: Optional[str] = None

class Education(BaseModel):
    school: str
    degree: str
    field: Optional[str] = None
    years: Optional[str] = None

class Certification(BaseModel):
    name: str
    issuer: str
    date: Optional[str] = None

class Volunteer(BaseModel):
    organization: str
    role: str
    duration: Optional[str] = None

class Recommendation(BaseModel):
    author: str
    relationship: str
    text: str

class LinkedInProfile(BaseModel):
    name: str
    headline: str
    location: str
    about: str
    experience: List[Experience]
    education: List[Education]
    skills: List[str]
    certifications: Optional[List[Certification]] = None
    languages: Optional[List[str]] = None
    volunteer: Optional[List[Volunteer]] = None
    recommendations: Optional[List[Recommendation]] = None

# Optional styling class
class ResumeStyle(BaseModel):
    """Defines the styling options for the DOCX resume (optional)"""
    font_name: str = "Calibri"
    name_size: int = 18
    heading_size: int = 14
    normal_size: int = 11
    heading_color: tuple = (0, 0, 0)  # RGB
    text_color: tuple = (0, 0, 0)     # RGB
    margins: float = 1.0              # inches
    line_spacing: float = 1.15

# ------------------------------------
# 2. GPT-based Structuring (Optional)
# ------------------------------------
def structure_profile_data(raw_text: str) -> LinkedInProfile:
    """
    Uses GPT-4o (example usage) to structure the raw LinkedIn profile text.
    Customize if you're not actually using GPT-4o or a similar approach.
    """
    client = OpenAI()
    
    completion = client.beta.chat.completions.parse(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": "Extract the LinkedIn profile information into a structured format."
            },
            {
                "role": "user",
                "content": raw_text
            }
        ],
        response_format=LinkedInProfile,
    )

    return completion.choices[0].message.parsed

def markdown_to_docx(markdown_file: str, output_file: str) -> str:
    """
    Converts the markdown resume to a simple DOCX file.
    """
    doc = Document()
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):  # Name
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):  # Headline
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):  # Section headers
            doc.add_heading(line[4:], 2)
        elif line.startswith('#### '):  # Subsections
            doc.add_heading(line[5:], 3)
        elif line.startswith('- '):  # List items
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('*') and line.endswith('*'):  # Italic text
            p = doc.add_paragraph()
            p.add_run(line.strip('*')).italic = True
        elif line.startswith('**') and line.endswith('**'):  # Bold text
            p = doc.add_paragraph()
            p.add_run(line.strip('**')).bold = True
        elif '**' in line:  # Handle inline bold text
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if part:  # Skip empty parts
                    run = p.add_run(part)
                    run.bold = (i % 2 == 1)  # Bold for odd-indexed parts
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(output_file)
    return output_file

def save_structured_profile(profile: LinkedInProfile, output_dir: str):
    """
    Saves the structured profile as markdown, HTML, and DOCX formats.
    """
    # ------------------------
    # Generate Markdown
    # ------------------------
    markdown = f"""# {profile.name}

## {profile.headline}
**Location:** {profile.location}

### About
{profile.about}

### Experience
"""
    
    for exp in profile.experience:
        markdown += f"""
#### {exp.title} at {exp.company}
*{exp.duration}*

{exp.description if exp.description else ''}
"""

    markdown += "\n### Education\n"
    for edu in profile.education:
        markdown += f"""
#### {edu.school}
{edu.degree}{f' in {edu.field}' if edu.field else ''}
{f'*{edu.years}*' if edu.years else ''}
"""

    if profile.skills:
        markdown += "\n### Skills\n"
        for skill in profile.skills:
            markdown += f"- {skill}\n"

    if profile.certifications:
        markdown += "\n### Certifications\n"
        for cert in profile.certifications:
            markdown += f"""
#### {cert.name}
Issued by {cert.issuer}{f' ({cert.date})' if cert.date else ''}
"""

    if profile.languages:
        markdown += "\n### Languages\n"
        for lang in profile.languages:
            markdown += f"- {lang}\n"

    if profile.recommendations:
        markdown += "\n### Recommendations\n"
        for rec in profile.recommendations:
            markdown += f"""
#### From {rec.author} ({rec.relationship})
{rec.text}
"""

    # ------------------------
    # Generate HTML (Optional)
    # ------------------------
    html = f"""
    <div class="profile-container">
        <header class="profile-header">
            <h1 class="profile-name">{profile.name}</h1>
            <h2 class="profile-headline">{profile.headline}</h2>
            <div class="profile-location">
                <i class="fas fa-map-marker-alt"></i> {profile.location}
            </div>
        </header>

        <section class="profile-section">
            <h3>About</h3>
            <div class="profile-about">
                {profile.about}
            </div>
        </section>

        <section class="profile-section">
            <h3>Experience</h3>
            <div class="experience-list">
                {''.join([f'''
                <div class="experience-item">
                    <div class="experience-header">
                        <h4>{exp.title}</h4>
                        <div class="company-name">{exp.company}</div>
                        <div class="duration">{exp.duration}</div>
                    </div>
                    <div class="experience-description">
                        {exp.description if exp.description else ''}
                    </div>
                </div>
                ''' for exp in profile.experience])}
            </div>
        </section>

        <section class="profile-section">
            <h3>Education</h3>
            <div class="education-list">
                {''.join([f'''
                <div class="education-item">
                    <div class="education-header">
                        <h4>{edu.school}</h4>
                        <div class="degree">
                            {edu.degree}{f' in {edu.field}' if edu.field else ''}
                        </div>
                        {f'<div class="years">{edu.years}</div>' if edu.years else ''}
                    </div>
                </div>
                ''' for edu in profile.education])}
            </div>
        </section>
    """

    if profile.skills:
        html += f"""
        <section class="profile-section">
            <h3>Skills</h3>
            <div class="skills-list">
                {''.join([f'<span class="skill-tag">{skill}</span>' for skill in profile.skills])}
            </div>
        </section>
        """

    if profile.certifications:
        html += f"""
        <section class="profile-section">
            <h3>Certifications</h3>
            <div class="certifications-list">
                {''.join([f'''
                <div class="certification-item">
                    <h4>{cert.name}</h4>
                    <div class="certification-meta">
                        <span class="issuer">Issued by {cert.issuer}</span>
                        {f'<span class="date">{cert.date}</span>' if cert.date else ''}
                    </div>
                </div>
                ''' for cert in profile.certifications])}
            </div>
        </section>
        """

    if profile.languages:
        html += f"""
        <section class="profile-section">
            <h3>Languages</h3>
            <div class="languages-list">
                {''.join([f'<span class="language-tag">{lang}</span>' for lang in profile.languages])}
            </div>
        </section>
        """

    if profile.recommendations:
        html += f"""
        <section class="profile-section recommendations-section">
            <h3>Recommendations</h3>
            <div class="recommendations-list">
                {''.join([f'''
                <div class="recommendation-item">
                    <div class="recommendation-header">
                        <div class="recommender">
                            <span class="recommender-name">{rec.author}</span>
                            <span class="relationship">({rec.relationship})</span>
                        </div>
                    </div>
                    <div class="recommendation-content">
                        "{rec.text}"
                    </div>
                </div>
                ''' for rec in profile.recommendations])}
            </div>
        </section>
        """

    html += "</div>"

    # ------------------------
    # Save All Files
    # ------------------------
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    markdown_file = os.path.join(output_dir, "structured_profile.md")
    with open(markdown_file, "w", encoding="utf-8") as f:
        f.write(markdown)
    
    html_file = os.path.join(output_dir, "structured_profile.html")
    with open(html_file, "w", encoding="utf-8") as f:
        f.write(html)
    
    docx_file = os.path.join(output_dir, "structured_profile.docx")
    markdown_to_docx(markdown_file, docx_file)
    
    return markdown_file, html_file, docx_file

# ------------------------------------
# 3. Main Extraction Logic
# ------------------------------------
def linkedin_highlight_and_extract(
    email: str,
    password: str,
    profile_url: str,
    output_dir="output"
):
    """
    Logs into LinkedIn with the provided credentials,
    navigates to the user-provided profile URL,
    expands hidden sections by clicking "see more" buttons,
    highlights everything on the page, then extracts all the text
    and saves it to 'profile.marathon'.
    
    Also includes an optional attempt to parse the data via GPT-4o,
    saving structured results if possible.
    """

    # ------------------------------
    # 1. Configure Selenium Options
    # ------------------------------
    chrome_options = Options()
    # (Optional) Uncomment if you want a headless browser:
    # chrome_options.add_argument('--headless')

    driver = webdriver.Chrome(options=chrome_options)

    try:
        # --------------------
        # 2. Log Into LinkedIn
        # --------------------
        driver.get("https://www.linkedin.com/login")

        # Wait for login form
        WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.ID, "username"))
        )

        email_input = driver.find_element(By.ID, "username")
        email_input.clear()
        email_input.send_keys(email)

        password_input = driver.find_element(By.ID, "password")
        password_input.clear()
        password_input.send_keys(password)

        # Submit form
        driver.find_element(By.XPATH, '//button[@type="submit"]').click()

        # ------------------------------------------------------
        # CAPTCHA / Verification Puzzle Handling (if it appears)
        # ------------------------------------------------------
        # This logic tries to detect a puzzle/captcha. If found, it pauses
        # and lets the user solve it manually in the browser, then continue.
        time.sleep(2)  # short wait for potential puzzle to appear

        # Example: we look for an element with class "captcha__prompt" or similar.
        # You may need to adjust the selector to match LinkedIn's actual puzzle/captcha.
        try:
            # If this doesn't exist, it will throw a TimeoutException
            puzzle_element = WebDriverWait(driver, 5).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, ".captcha__prompt, .rc-imageselect-tile"))  # Example CSS
            )
            if puzzle_element:
                print("\nA verification puzzle has appeared.")
                print("Please solve it in the browser window. Then press Enter here to continue.")
                input("Press Enter once the puzzle is solved...")
        except:
            # If no puzzle is found within 5s, we assume no captcha is needed
            pass

        # ------------------------------------------------
        # 3. Navigate to the user-specified profile URL
        # ------------------------------------------------
        WebDriverWait(driver, 30).until(
            EC.presence_of_element_located((By.TAG_NAME, "body"))
        )

        time.sleep(3)  # Additional wait after login
        driver.get(profile_url)

        time.sleep(5)  # Wait for the profile page to load

        # -------------------------------------------------------
        # 4. Locate & Click "see more" Buttons to Expand Sections
        # -------------------------------------------------------
        see_more_selectors = [
            ".inline-show-more-text__button.inline-show-more-text__button--light.link"
        ]
        for selector in see_more_selectors:
            buttons = driver.find_elements(By.CSS_SELECTOR, selector)
            for btn in buttons:
                try:
                    driver.execute_script("arguments[0].click();", btn)
                    time.sleep(1)
                except Exception as e:
                    print(f"Warning: Could not click a 'see more' button: {e}")

        # -------------------------
        # 5. "Highlight Everything"
        # -------------------------
        # Simulate Ctrl + A in the browser
        driver.execute_script("window.getSelection().removeAllRanges();")
        driver.execute_script("const range = document.createRange(); range.selectNode(document.body); window.getSelection().addRange(range);")

        # -------------------------------
        # 6. Extract All Visible Text
        # -------------------------------
        page_text = driver.execute_script("return document.body.innerText")

        # ---------------------------------
        # 7. Save the text into .marathon
        # ---------------------------------
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        marathon_file = os.path.join(output_dir, "profile.marathon")

        with open(marathon_file, "w", encoding="utf-8") as f:
            f.write(page_text)

        print(f"Entire expanded profile text saved to: {marathon_file}")

        # ------------------------------------------------
        # 8. (Optional) Parse & Save Structured Versions
        # ------------------------------------------------
        try:
            structured_profile = structure_profile_data(page_text)
            
            raw_file = os.path.join(output_dir, "profile.marathon")
            with open(raw_file, "w", encoding="utf-8") as f:
                f.write(page_text)
            
            # Save structured as MD/HTML/DOCX
            structured_files = save_structured_profile(structured_profile, output_dir)

            print(f"\nStructured profile saved. Files generated:")
            print(f"1) {structured_files[0]} (Markdown)")
            print(f"2) {structured_files[1]} (HTML)")
            print(f"3) {structured_files[2]} (DOCX)")

            return structured_profile

        finally:
            driver.quit()

    except Exception as e:
        print(f"Error: {e}")
        driver.quit()

# --------------------------------
# 4. Main Entry Point
# --------------------------------
def main():
    print("LinkedIn Highlight & Extract")
    print("-------------------------------------")
    email = input("LinkedIn Email/Username: ").strip()
    password = getpass.getpass("LinkedIn Password: ")
    profile_url = input("LinkedIn Profile URL (e.g., https://www.linkedin.com/in/username/): ").strip()

    profile = linkedin_highlight_and_extract(email, password, profile_url)
    
    if profile:
        print("\nProfile extracted successfully!")
        print("You can find your files in the output directory:")
        print("1. profile.marathon (raw text)")
        print("2. structured_profile.md (if GPT parsing succeeds)")
        print("3. structured_profile.docx (Word doc)")
        print("4. structured_profile.html (HTML version)")

if __name__ == "__main__":
    main()
